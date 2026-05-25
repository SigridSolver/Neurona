import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Colors
    COLOR_PRIMARY = RGBColor(30, 41, 59)      # Slate 800
    COLOR_SECONDARY = RGBColor(59, 130, 246)   # Blue 500
    COLOR_TEXT = RGBColor(71, 85, 105)        # Slate 600
    
    # Helper for heading styling
    def style_heading(text, level, space_before=18, space_after=6):
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        
        # Color and Font
        for run in heading.runs:
            run.font.name = 'Arial'
            if level == 1:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = COLOR_PRIMARY
            elif level == 2:
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = COLOR_SECONDARY
        return heading

    # Helper for paragraphs
    def add_p(text="", bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_bold = p.add_run(bold_prefix)
            run_bold.font.name = 'Arial'
            run_bold.font.size = Pt(11)
            run_bold.font.bold = True
            run_bold.font.color.rgb = COLOR_PRIMARY
            
        if text:
            run_text = p.add_run(text)
            run_text.font.name = 'Arial'
            run_text.font.size = Pt(11)
            run_text.font.color.rgb = COLOR_TEXT
        return p

    # Callout Box / Alert Style
    def add_callout(text, title="NOTA"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell = table.cell(0, 0)
        
        # Set shading (background light blue)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F6FF"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
        # Set left border (thick blue)
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="3B82F6"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        
        run_title = p.add_run(f"💡 {title}: ")
        run_title.font.name = 'Arial'
        run_title.font.size = Pt(10.5)
        run_title.font.bold = True
        run_title.font.color.rgb = COLOR_SECONDARY
        
        run_text = p.add_run(text)
        run_text.font.name = 'Arial'
        run_text.font.size = Pt(10.5)
        run_text.font.italic = True
        run_text.font.color.rgb = COLOR_PRIMARY
        
        # Add space after table
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(6)

    # Document Header / Title Page Info
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)
    run_t = title_p.add_run("INFORME PROYECTO PEDAGÓGICO")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    run_t.font.color.rgb = COLOR_SECONDARY

    main_title = doc.add_paragraph()
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title.paragraph_format.space_before = Pt(4)
    main_title.paragraph_format.space_after = Pt(24)
    run_mt = main_title.add_run("NEURONA SABER 11\nPlataforma Educativa Gamificada con Agente de Inteligencia Artificial")
    run_mt.font.name = 'Arial'
    run_mt.font.size = Pt(20)
    run_mt.font.bold = True
    run_mt.font.color.rgb = COLOR_PRIMARY

    # Metadata metadata
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(36)
    run_meta = meta_p.add_run("Presentado a: Director de Curso\nProyecto de Innovación y Preparación Académica\nColegio - 2026")
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(11)
    run_meta.font.color.rgb = COLOR_TEXT

    # Section 1
    style_heading("1. Descripción del Proyecto", level=1)
    add_p("es una plataforma web educativa innovadora diseñada específicamente para potenciar la preparación de los estudiantes de bachillerato de cara a las pruebas de estado Saber 11 (ICFES). A diferencia de los métodos tradicionales de estudio pasivos o memorísticos, Neurona integra técnicas avanzadas de gamificación, interacción competitiva entre pares y un Agente de Inteligencia Artificial para el acompañamiento adaptativo en tiempo real.", bold_prefix="Neurona Saber 11 ")
    
    add_p("El objetivo principal es estructurar un ecosistema de aprendizaje autónomo y motivador, donde cada estudiante reciba recomendaciones precisas, entienda de manera inmediata sus errores conceptuales y sea incentivado diariamente a través de mecánicas de superación colectiva.")

    # Section 2
    style_heading("2. Secciones Principales de la Aplicación", level=1)
    
    style_heading("A. Panel de Control Personalizado (Dashboard)", level=2)
    add_p("Actúa como el centro neurálgico del estudiante. En esta sección se visualiza el progreso global acumulado, la racha diaria de estudio, los puntos obtenidos y el nivel del usuario. Asimismo, integra un radar interactivo en forma de pentágono que mapea las fortalezas y áreas de mejora del alumno en las cinco áreas evaluadas por el ICFES.")
    
    style_heading("B. Módulo de Práctica Temática y Libre", level=2)
    add_p("Permite entrenarse libremente seleccionando el área de interés (Matemáticas, Lectura Crítica, Ciencias Naturales, Sociales y Ciudadanas, o Inglés) y el nivel de dificultad. Tras responder cada pregunta, la plataforma ofrece retroalimentación inmediata, resaltando la respuesta correcta y mostrando la explicación teórica provista por el Agente de IA para disipar cualquier vacío didáctico al instante.")

    style_heading("C. Simulacros y Diagnóstico General", level=2)
    add_p("Una recreación rigurosa del examen real con control de tiempo. Sirve como prueba diagnóstica para calcular el desempeño esperado bajo condiciones similares a la prueba real y alimenta los modelos del Agente Pedagógico para habilitar o condicionar funcionalidades adicionales del sistema.")

    style_heading("D. Sala de Duelos 1v1 (Desafíos de Estudio)", level=2)
    add_p("Espacio competitivo donde un estudiante puede retar a un compañero o responder retos entrantes en tiempo real o asíncronamente en un lapso de 24 horas. Ambos responden la misma pregunta y el sistema evalúa los aciertos, fomentando una dinámica cooperativa de sana competencia en la institución educativa.")

    # Section 3
    style_heading("3. Mecánica de Gamificación, Puntos y Rankings", level=1)
    add_p("Para mantener altos niveles de participación activa, el sistema implementa una lógica de recompensas gamificadas basada en méritos y participación:")
    
    # Bullet points
    p_b1 = doc.add_paragraph(style='List Bullet')
    p_b1.paragraph_format.space_after = Pt(3)
    p_b1.add_run("Victoria en Duelos: ").bold = True
    p_b1.runs[0].font.name = 'Arial'
    p_b1.add_run("El estudiante que gana el desafío es premiado con 5 puntos de experiencia en el ranking general.").font.name = 'Arial'
    
    p_b2 = doc.add_paragraph(style='List Bullet')
    p_b2.paragraph_format.space_after = Pt(3)
    p_b2.add_run("Incentivo al Desafiar (Retador): ").bold = True
    p_b2.runs[0].font.name = 'Arial'
    p_b2.add_run("Con el fin de fomentar la iniciativa y participación activa, se otorga 1 punto de incentivo al estudiante que envía el duelo (retador), independientemente de si gana, pierde o empata al final.").font.name = 'Arial'

    p_b3 = doc.add_paragraph(style='List Bullet')
    p_b3.paragraph_format.space_after = Pt(3)
    p_b3.add_run("Empate o Derrota: ").bold = True
    p_b3.runs[0].font.name = 'Arial'
    p_b3.add_run("Un duelo empatado o perdido otorga 0 puntos de experiencia a las respuestas respectivas (manteniendo únicamente el punto de incentivo si se es el emisor del reto).").font.name = 'Arial'

    p_b4 = doc.add_paragraph(style='List Bullet')
    p_b4.paragraph_format.space_after = Pt(6)
    p_b4.add_run("Ranking y Tabla de Líderes: ").bold = True
    p_b4.runs[0].font.name = 'Arial'
    p_b4.add_run("Clasificación visible en la sección de duelos que muestra en tiempo real las victorias acumuladas, los puntos totales de experiencia, los niveles alcanzados y medallas institucionales obtenidas por cada estudiante.").font.name = 'Arial'

    # Section 4
    style_heading("4. Integración del Agente de Inteligencia Artificial Pedagógica", level=1)
    add_p("El proyecto cuenta con un Agente AI enfocado en el contexto educativo que actúa como un tutor virtual permanente y adaptativo. Sus dos funciones principales son:")
    
    add_p("El agente genera automáticamente análisis lógicos estructurados para cada pregunta fallada u omitida, explicando de forma clara y constructiva por qué cada distractor es incorrecto y justificando rigurosamente la opción verdadera.", bold_prefix="Retroalimentación Explicativa: ")
    
    add_p("Un sistema inteligente de tutoría activado mediante el icono del bombillo (💡) que ofrece consejos y pistas breves sin revelar directamente la respuesta. Sin embargo, para fomentar la disciplina de auto-evaluación inicial, esta función está sujeta a restricciones académicas:", bold_prefix="Sistema de Pistas y Ayudas (Tips de Estudio): ")

    # Add Callout box for bulb rules
    add_callout(
        "Las ayudas pedagógicas del Agente AI solo se desbloquean si el estudiante ha realizado previamente su Simulacro de Diagnóstico y ha superado un umbral de 250 puntos globales (sobre 500). Para los estudiantes con menor puntaje, el bombillo 💡 permanece visible pero atenuado; al hacer clic, el Agente les presenta un mensaje motivacional invitándoles a concentrarse en los simulacros e incrementar su puntuación para desbloquear el módulo avanzado de tutoría.",
        "Condición de Acceso por Mérito Académico"
    )

    # Section 5
    style_heading("5. Impacto Educativo del Proyecto", level=1)
    add_p("La implementación de Neurona Saber 11 en el aula de clases genera múltiples impactos positivos:")
    
    # Bullets
    p_i1 = doc.add_paragraph(style='List Bullet')
    p_i1.paragraph_format.space_after = Pt(3)
    p_i1.add_run("Mitigación de la Ansiedad al Examen: ").bold = True
    p_i1.runs[0].font.name = 'Arial'
    p_i1.add_run("Al convertir la preparación en una actividad interactiva y cotidiana con mecánicas lúdicas, disminuye el estrés del estudiante ante el ICFES.").font.name = 'Arial'

    p_i2 = doc.add_paragraph(style='List Bullet')
    p_i2.paragraph_format.space_after = Pt(3)
    p_i2.add_run("Acompañamiento 24/7 sin Sobrecarga Docente: ").bold = True
    p_i2.runs[0].font.name = 'Arial'
    p_i2.add_run("El Agente de IA resuelve dudas teóricas inmediatas a cualquier hora, liberando al docente para tutorías grupales de alto valor.").font.name = 'Arial'

    p_i3 = doc.add_paragraph(style='List Bullet')
    p_i3.paragraph_format.space_after = Pt(6)
    p_i3.add_run("Banco de Preguntas Institucional: ").bold = True
    p_i3.runs[0].font.name = 'Arial'
    p_i3.add_run("La base de datos local almacena y clasifica preguntas reales y simuladas de nivel institucional, permitiendo al colegio contar con un banco digital auditable de alto nivel académico.").font.name = 'Arial'

    # Section 6
    style_heading("6. Conclusiones del Proyecto", level=1)
    add_p("El proyecto representa un cambio de paradigma en la forma en que los estudiantes abordan la preparación de pruebas estandarizadas de alta relevancia en Colombia. La fusión de gamificación y acompañamiento inteligente no solo optimiza el rendimiento académico, sino que también fomenta hábitos de estudio saludables, regulares y colaborativos.")
    add_p("Al condicionar las ayudas avanzadas (💡) a un rendimiento mínimo en simulacros, se introduce un incentivo directo al esfuerzo sostenido, estimulando al estudiante a mejorar progresivamente y a utilizar los simulacros como verdaderos termómetros de su preparación.")

    # Section 7
    style_heading("7. Futuras Mejoras y Evolución del Sistema", level=1)
    add_p("Para alinearse con las directrices de evaluación institucional y las políticas de mejoramiento continuo del ICFES, se contemplan las siguientes líneas de evolución y escalado para la plataforma:")

    p_f1 = doc.add_paragraph(style='List Bullet')
    p_f1.paragraph_format.space_after = Pt(3)
    p_f1.add_run("Módulo de Gestión Docente (Panel de Tutores): ").bold = True
    p_f1.runs[0].font.name = 'Arial'
    p_f1.add_run("Incorporación de un entorno administrativo para que los profesores puedan realizar un seguimiento focalizado de los alumnos, diseñar cuestionarios personalizados, asignar tareas libres o exámenes específicos para un grupo, y analizar reportes grupales e individuales de debilidades de aprendizaje.").font.name = 'Arial'

    p_f2 = doc.add_paragraph(style='List Bullet')
    p_f2.paragraph_format.space_after = Pt(3)
    p_f2.add_run("Diversificación de Pruebas de Estado (Saber Multigrado): ").bold = True
    p_f2.runs[0].font.name = 'Arial'
    p_f2.add_run("Ampliación del banco de preguntas y adaptabilidad del sistema para cubrir todo el espectro de evaluación del ICFES, incluyendo pruebas de educación básica (Saber 3°, Saber 5°, Saber 9°) y educación superior (Saber Pro y Saber TyT), estructurando un historial de aprendizaje transversal.").font.name = 'Arial'

    p_f3 = doc.add_paragraph(style='List Bullet')
    p_f3.paragraph_format.space_after = Pt(3)
    p_f3.add_run("Nuevas Áreas Interactivas y Cultura General: ").bold = True
    p_f3.runs[0].font.name = 'Arial'
    p_f3.add_run("Integración de módulos adicionales para evaluar cultura general, competencias blandas, quizzes en áreas artísticas, tecnológicas y éticas, ampliando el impacto pedagógico de la plataforma más allá del currículo evaluativo tradicional.").font.name = 'Arial'

    # Save
    filename = "Resumen_Ejecutivo_Proyecto_Saber11.docx"
    try:
        doc.save(filename)
        print(f"Documento guardado exitosamente como {filename}")
    except PermissionError:
        alt_filename = "Resumen_Ejecutivo_Proyecto_Saber11_Actualizado.docx"
        doc.save(alt_filename)
        print(f"Documento guardado exitosamente como {alt_filename} (el original estaba abierto).")

if __name__ == '__main__':
    create_report()
